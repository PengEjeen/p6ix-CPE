"""Word report document builder."""

from datetime import date as date_cls
from io import BytesIO

from .cs_common import apply_center_page_number_footer, setup_document_defaults
from .cs_section_duration import add_duration_analysis_section


def _set_section_a4(section, landscape=False):
    from docx.enum.section import WD_ORIENTATION
    from docx.shared import Mm

    side_margin = Mm(10)

    if landscape:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
    else:
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)

    section.left_margin = side_margin
    section.right_margin = side_margin


def _force_continuous_page_numbering(document):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for section in document.sections:
        sect_pr = section._sectPr
        for child in list(sect_pr):
            if child.tag == qn("w:pgNumType"):
                sect_pr.remove(child)

    for idx, section in enumerate(document.sections):
        sect_pr = section._sectPr
        pg_num = OxmlElement("w:pgNumType")
        pg_num.set(qn("w:fmt"), "decimal")
        # Important: only the first section gets explicit start=1.
        # Later sections must NOT inherit this start value, otherwise
        # PAGEREF/PAGE can appear to restart at 1 per section.
        if idx == 0:
            pg_num.set(qn("w:start"), "1")
        sect_pr.append(pg_num)


def _add_cover_page(document, project_name):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    project_title = str(project_name or "프로젝트").strip() or "프로젝트"
    report_title = "[ 공기적정성 검토 보고서 ]"
    today_text = date_cls.today().strftime("%Y. %m.")

    def _blank_lines(count):
        for _ in range(count):
            p = document.add_paragraph("")
            p.paragraph_format.space_after = Pt(0)

    _blank_lines(9)

    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(project_title)
    run.bold = True
    run.font.size = Pt(24)

    p_subtitle = document.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subtitle.paragraph_format.space_before = Pt(14)
    run = p_subtitle.add_run(report_title)
    run.bold = True
    run.font.size = Pt(20)

    _blank_lines(17)

    p_date = document.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_date.add_run(today_text)
    run.bold = True
    run.font.size = Pt(16)

    document.add_page_break()


def _add_toc_page(document, weather_appendix_data=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def _append_pageref_field(paragraph, bookmark_name, font_size_pt):
        # Use fldSimple for better compatibility in Word field parsing.
        fld_simple = OxmlElement("w:fldSimple")
        fld_simple.set(qn("w:instr"), f"PAGEREF {bookmark_name}")
        fld_simple.set(qn("w:dirty"), "true")

        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        font_half_pt = str(int(round(float(font_size_pt.pt) * 2)))
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), font_half_pt)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), font_half_pt)
        rpr.append(sz)
        rpr.append(sz_cs)
        run.append(rpr)

        text = OxmlElement("w:t")
        text.text = "1"
        run.append(text)
        fld_simple.append(run)
        paragraph._p.append(fld_simple)

    p_heading = document.add_paragraph()
    p_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_heading.add_run("목차")
    run.bold = True
    run.font.size = Pt(22)
    document.add_paragraph("")

    section = document.sections[-1]
    content_width = section.page_width - section.left_margin - section.right_margin

    toc_rows = [
        (0, "제1장 공사기간 분석", "bm_chapter_1"),
        (1, "1.1 공사기간 산정 (국토교통부 공사기간 산정 기준)", "bm_sec_1_1"),
        (1, "1.2 근로시간 적용 기준", "bm_sec_1_2"),
        (1, "1.3 비작업일수 산정", "bm_sec_1_3"),
        (1, "1.4 공종별 표준작업량 산정", "bm_sec_1_4"),
        (1, "1.5 실공사기간 Calendar Day 산출", "bm_sec_1_5"),
        (1, "1.6 공사예정공정표 작성기준", "bm_sec_1_6"),
        (1, "1.7 공사예정 공정표", "bm_sec_1_7"),
        (1, "1.8 별첨", "bm_sec_1_8"),
    ]

    appendix_data = weather_appendix_data or {}
    for appendix in (appendix_data.get("appendices") or []):
        number = appendix.get("number")
        title = appendix.get("title", "")
        toc_rows.append((2, f"별첨.{number} {title}", f"bm_appendix_{number}"))

    yearly_status = appendix_data.get("yearly_status") or {}
    if yearly_status.get("tables"):
        toc_rows.append((2, "<별첨> 년도별 기상 휴지일수 현황", "bm_appendix_yearly_status"))

    for level, title, bookmark_name in toc_rows:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Pt(12 * level)
        p.paragraph_format.space_after = Pt(5 if level == 0 else 3)
        p.paragraph_format.tab_stops.add_tab_stop(
            content_width,
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )

        text_run = p.add_run(str(title))
        text_run.font.size = Pt(13 if level == 0 else 11)
        text_run.bold = (level == 0)

        p.add_run("\t")
        _append_pageref_field(
            p,
            bookmark_name,
            Pt(13 if level == 0 else 11),
        )

def build_schedule_report_docx(
    project_name,
    rate_summary,
    ordered_categories,
    grouped_items,
    rate_map,
    region="",
    gantt_image_bytes=None,
    project_overview=None,
    public_holiday_rows=None,
    climate_criteria_rows=None,
    climate_category_rows=None,
    monthly_condition_rows=None,
    monthly_year_range="",
    operating_rate_calc_data=None,
    weather_appendix_data=None,
):
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for report export. Install 'python-docx' in backend requirements."
        ) from exc

    # Keep signature for compatibility with current view calls.
    _ = (project_name, rate_summary, ordered_categories, grouped_items, rate_map, region)

    document = Document()
    setup_document_defaults(document)

    # Ask Word to refresh fields (PAGEREF/TOC-like fields) when opening.
    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    _set_section_a4(document.sections[0], landscape=False)
    _add_cover_page(document, project_name=project_name)
    _add_toc_page(document, weather_appendix_data=weather_appendix_data)

    from docx.enum.section import WD_SECTION_START
    body_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    _set_section_a4(body_section, landscape=False)

    add_duration_analysis_section(
        document,
        ordered_categories=ordered_categories,
        grouped_items=grouped_items,
        rate_map=rate_map,
        project_name=project_name,
        gantt_image_bytes=gantt_image_bytes,
        project_overview=project_overview,
        public_holiday_rows=public_holiday_rows,
        climate_criteria_rows=climate_criteria_rows,
        climate_category_rows=climate_category_rows,
        monthly_condition_rows=monthly_condition_rows,
        monthly_year_range=monthly_year_range,
        operating_rate_calc_data=operating_rate_calc_data,
        weather_appendix_data=weather_appendix_data,
    )

    _force_continuous_page_numbering(document)
    apply_center_page_number_footer(document, skip_cover_page=True)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.read()
