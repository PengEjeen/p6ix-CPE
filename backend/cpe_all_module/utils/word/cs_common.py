"""Common helpers for Word report rendering."""

from decimal import Decimal


def set_cell_shading(cell, fill="EDEDED"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:shd"):
            tc_pr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def shade_header_row(table, row_idx=0, fill="BFBFBF"):
    if row_idx < 0 or row_idx >= len(table.rows):
        return
    for cell in table.rows[row_idx].cells:
        set_cell_shading(cell, fill)


def merge_same_text_cells(table, col_idx, start_row=1):
    row_count = len(table.rows)
    row = start_row
    while row < row_count:
        current_text = (table.cell(row, col_idx).text or "").strip()
        end = row
        while end + 1 < row_count:
            next_text = (table.cell(end + 1, col_idx).text or "").strip()
            if next_text != current_text:
                break
            end += 1
        if end > row and current_text:
            merged_cell = table.cell(row, col_idx).merge(table.cell(end, col_idx))
            merged_cell.text = current_text
        row = end + 1


def to_number(value):
    if value is None or value == "":
        return None
    if isinstance(value, (int, float, Decimal)):
        return float(value)
    try:
        return float(str(value).replace(",", ""))
    except (TypeError, ValueError):
        return None


def format_number(value, digits=1):
    number = to_number(value)
    if number is None:
        return "-"
    if digits == 0:
        return f"{number:,.0f}"
    return f"{number:,.{digits}f}"


def setup_document_defaults(document):
    from docx.shared import Pt

    style = document.styles["Normal"]
    style.font.size = Pt(10)
    try:
        from docx.enum.section import WD_ORIENTATION
    except ImportError:
        return
    for section in document.sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width


def apply_center_page_number_footer(document, skip_cover_page=False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _append_page_field(paragraph):
        fld_simple = OxmlElement("w:fldSimple")
        fld_simple.set(qn("w:instr"), "PAGE")
        fld_simple.set(qn("w:dirty"), "true")

        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "1"
        run.append(text)
        fld_simple.append(run)
        paragraph._p.append(fld_simple)

    for section_idx, section in enumerate(document.sections):
        section.footer.is_linked_to_previous = False
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()

        # Keep only one footer paragraph to avoid style/tab side effects.
        for extra in footer.paragraphs[1:]:
            extra._element.getparent().remove(extra._element)

        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0

        # Keep paragraph properties (w:pPr), remove only content runs/fields.
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)

        paragraph.add_run(" - ")
        _append_page_field(paragraph)
        paragraph.add_run(" - ")

        # Hide page number only on the very first page (cover page).
        if skip_cover_page and section_idx == 0:
            section.different_first_page_header_footer = True
            first_footer = section.first_page_footer
            if not first_footer.paragraphs:
                first_footer.add_paragraph()
            for extra in first_footer.paragraphs[1:]:
                extra._element.getparent().remove(extra._element)
            first_paragraph = first_footer.paragraphs[0]
            first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for child in list(first_paragraph._p):
                first_paragraph._p.remove(child)
