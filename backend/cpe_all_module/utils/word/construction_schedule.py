from io import BytesIO
from decimal import Decimal


def _set_cell_shading(cell, fill="EDEDED"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    # Replace existing shading if present.
    for child in list(tc_pr):
        if child.tag == qn("w:shd"):
            tc_pr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _shade_header_row(table, row_idx=0, fill="BFBFBF"):
    if row_idx < 0 or row_idx >= len(table.rows):
        return
    for cell in table.rows[row_idx].cells:
        _set_cell_shading(cell, fill)


def _merge_same_text_cells(table, col_idx, start_row=1):
    row_count = len(table.rows)
    row = start_row
    while row < row_count:
        current_text = (table.cell(row, col_idx).text or "").strip()
        end = row
        while end + 1 < row_count:
            next_text = (table.cell(end + 1, col_idx).text or "").strip()
            if next_text != current_text:
                break
            end += 1
        if end > row and current_text:
            merged_cell = table.cell(row, col_idx).merge(table.cell(end, col_idx))
            # Keep only one label text after merge (avoid duplicated concatenation).
            merged_cell.text = current_text
        row = end + 1


def _to_number(value):
    if value is None or value == "":
        return None
    if isinstance(value, (int, float, Decimal)):
        return float(value)
    try:
        return float(str(value).replace(",", ""))
    except (TypeError, ValueError):
        return None


def _format_number(value, digits=1):
    number = _to_number(value)
    if number is None:
        return "-"
    if digits == 0:
        return f"{number:,.0f}"
    return f"{number:,.{digits}f}"


def _setup_document_defaults(document):
    # Keep a neutral default style. Korean glyphs follow system fallback font.
    style = document.styles["Normal"]
    style.font.size = document.styles["Normal"].font.size


def _add_duration_analysis_section(document, public_holiday_rows=None, climate_criteria_rows=None):
    document.add_heading("1. 공사기간 분석", level=1)
    document.add_heading("1.1 공사기간 산정 (국토교통부 공사기간 산정 기준)", level=2)
    document.add_heading("1.1.1 공사기간의 정의", level=3)

    document.add_paragraph(
        "• 공사기간은 건설공사 계약의 착수일로부터 완료일까지의 기간을 의미하며, "
        "준비기간과 비작업일수, 작업일수, 정리기간을 포함하여 산정한다."
    )

    p_formula = document.add_paragraph()
    p_formula.add_run("공사기간 = ①준비기간 + ②비작업일수 + ③작업일수 + ④정리기간").bold = True

    document.add_paragraph(
        "① 준비기간: 설계도서 검토, 안전관리계획서 작성·승인, 하도급업체 선정, 측량, "
        "현장사무소·세륜시설·가설건물 설치, 주요 자재·장비 조달 등 공사 착공 준비에 필요한 기간"
    )
    document.add_paragraph(
        "② 비작업일수: 법정공휴일 및 기상조건 등으로 작업이 불가능한 일수를 고려하여 산정한 기간"
    )
    document.add_paragraph(
        "③ 작업일수: 해당 공사의 공종별 수량을 실제 시공하는 데 필요한 총 작업일수"
    )
    document.add_paragraph(
        "④ 정리기간: 준공검사 준비, 시설물 인수 및 청소 등 현장 정리에 필요한 기간"
    )

    document.add_heading("1.1.2 ①준비기간과 ④정리기간", level=3)
    document.add_paragraph("• 준비기간 및 정리기간은 비작업일수를 별도로 계산하지 않는다.")
    document.add_paragraph(
        "• 공사 유형별 준비기간 예시를 참고하고 본 공사의 특성을 고려하여 산정한다 "
        "(예: 준비기간 45일 적용)."
    )

    document.add_paragraph("<참고> 건설공사 유형별 준비기간")
    ref_rows = [
        ("공동주택", "45일", "상수도공사", "60일"),
        ("고속도로공사", "180일", "하천공사", "40일"),
        ("철도공사", "90일", "항만공사", "40일"),
        ("포장공사(신설)", "50일", "강교가설공사", "90일"),
        ("포장공사(수선)", "60일", "PC교량공사", "70일"),
        ("공동구공사", "80일", "교량보수공사", "60일"),
    ]
    ref_table = document.add_table(rows=1, cols=4)
    ref_table.style = "Table Grid"
    ref_headers = ref_table.rows[0].cells
    ref_headers[0].text = "공종"
    ref_headers[1].text = "준비기간"
    ref_headers[2].text = "공종"
    ref_headers[3].text = "준비기간"
    _shade_header_row(ref_table)
    for left_type, left_days, right_type, right_days in ref_rows:
        row = ref_table.add_row().cells
        row[0].text = left_type
        row[1].text = left_days
        row[2].text = right_type
        row[3].text = right_days

    document.add_paragraph(
        "• 정리기간은 공정상 여유기간(buffer)과는 다르며, 공사 규모 및 난이도 등을 고려하여 산정한다."
    )
    document.add_paragraph(
        "• 일반적으로 주요 공종이 마무리된 이후 준공 전 1개월 범위에서 계산하며 "
        "(예: 정리기간 30일 적용)."
    )

    document.add_heading("1.2 근로시간 적용 기준", level=2)
    document.add_heading("1.2.1 근로시간 정의", level=3)
    document.add_paragraph("• 2021년 개정된 「근로기준법」상의 기준시간의 정의는 다음과 같다.")

    document.add_paragraph("<참고> 근로기준법 제50조 근로시간")
    labor50_table = document.add_table(rows=2, cols=2)
    labor50_table.style = "Table Grid"
    labor50_table.rows[0].cells[0].text = "조항"
    labor50_table.rows[0].cells[1].text = "조문"
    _shade_header_row(labor50_table)
    labor50_table.rows[1].cells[0].text = "제50조\n(근로시간)"
    labor50_table.rows[1].cells[1].text = (
        "① 1주 간의 근로시간은 휴게시간을 제외하고 40시간을 초과할 수 없다.\n"
        "② 1일의 근로시간은 휴게시간을 제외하고 8시간을 초과할 수 없다.\n"
        "③ 제1항 및 제2항에 따라 근로시간을 산정하는 경우 작업을 위하여 "
        "근로자가 사용자의 지휘·감독 아래에 있는 대기시간 등은 근로시간으로 본다.\n"
        "<신설 2012. 2. 1., 2020. 5. 26.>"
    )

    document.add_paragraph(
        "• 국토교통부 고시 「공공 건설공사의 공사기간 산정기준」상의 기준시간은 다음과 같다."
    )
    document.add_paragraph("<참고> 국토교통부 고시 공공 건설공사의 공사기간 산정기준 제11조 3항")
    guideline_table = document.add_table(rows=2, cols=2)
    guideline_table.style = "Table Grid"
    guideline_table.rows[0].cells[0].text = "조항"
    guideline_table.rows[0].cells[1].text = "조문"
    _shade_header_row(guideline_table)
    guideline_table.rows[1].cells[0].text = "제11조\n(작업일수)"
    guideline_table.rows[1].cells[1].text = (
        "③ 작업일수 산정 시 건설현장 근로자의 작업조건이 법정 근로시간"
        "(1일 8시간, 주 40시간)을 준수하는 것을 원칙으로 한다. 다만, 연속작업 등에 "
        "필요한 경우에는 근로기준법에 따라 근로시간을 연장할 수 있고, 과다근무 및 "
        "주·야간 공사를 구분하여 산출한다."
    )

    document.add_paragraph(
        "• 상기 법령의 정의를 기준으로 작업일수 산정 시 반영 → "
        "주 6일(일 8시간 작업) 적용"
    )
    document.add_paragraph(
        "• 단, 터널은 레미콘 주5일제/시험 발파 → "
        "주 5일(일 8시간 작업) 적용"
    )

    document.add_paragraph("<참고> 근로기준법 제53조 연장 근로의 제한")
    labor53_table = document.add_table(rows=2, cols=2)
    labor53_table.style = "Table Grid"
    labor53_table.rows[0].cells[0].text = "조항"
    labor53_table.rows[0].cells[1].text = "조문"
    _shade_header_row(labor53_table)
    labor53_table.rows[1].cells[0].text = "제53조\n(연장 근로의 제한)"
    labor53_table.rows[1].cells[1].text = (
        "② 당사자 간에 합의하면 1주 간에 12시간을 한도로 제51조 및 제51조의2의 "
        "근로시간을 연장할 수 있고, 제52조 제1항 제2호의 시간간주를 평균하여 "
        "1주 간에 12시간을 초과하지 아니하는 범위에서 제52조 제1항의 근로시간을 "
        "연장할 수 있다.\n"
        "<개정 2021. 1. 5.>"
    )

    document.add_paragraph("• 상기 개정 법령의 연장근로는 주당 최대 12시간 그대로 유지")
    document.add_heading("1.2.2 근로시간 적용", level=3)
    document.add_paragraph(
        "• 근로기준법 및 국토교통부 고시의 기준을 바탕으로 개정 전/후를 아래와 같이 비교한다."
    )

    document.add_paragraph("<개정 전·후 주 최대 근로시간 비교표>")
    compare_table = document.add_table(rows=3, cols=5)
    compare_table.style = "Table Grid"
    compare_table.rows[0].cells[0].text = "기준(개정 전)"
    compare_table.rows[0].cells[1].text = "법정근로\n주 40시간"
    compare_table.rows[0].cells[2].text = "휴일근로\n16시간"
    compare_table.rows[0].cells[3].text = "주 최대 근로시간\n56시간"
    compare_table.rows[0].cells[4].text = "연장근로량\n12시간"
    _shade_header_row(compare_table)

    compare_table.rows[1].cells[0].text = "개정(2018.7.1.)"
    compare_table.rows[1].cells[1].text = "법정근로\n주 40시간"
    compare_table.rows[1].cells[2].text = "-"
    compare_table.rows[1].cells[3].text = "주 최대 근로시간\n40시간"
    compare_table.rows[1].cells[4].text = "-"

    compare_table.rows[2].cells[0].text = "변경사항"
    compare_table.rows[2].cells[1].text = "<상동>"
    compare_table.rows[2].cells[2].text = "<감소>\n<16시간 감소>"
    compare_table.rows[2].cells[3].text = "<감소>"
    compare_table.rows[2].cells[4].text = "<상동>"

    document.add_paragraph(
        "• 기존 근로시간과의 차이로 휴일근로 16시간 감소에 따라 연장근로 12시간은 "
        "시공사가 필요 시 법령에 따라 실시하며, 근로시간은 간주하여 공기산정 "
        "근로시간에서는 제외한다."
    )
    document.add_paragraph("• 근로시간 적용방안의 예시는 아래와 같다.")

    document.add_paragraph("<참고> 근로시간 적용 방안 예시")
    policy_table = document.add_table(rows=2, cols=3)
    policy_table.style = "Table Grid"
    policy_table.rows[0].cells[0].text = "법정 근로시간(주 5일, 40시간)에 따른\n1달 공휴일 수 8일"
    policy_table.rows[0].cells[1].text = "기후여건으로 인한\n불능일 < 8일 적용 경우"
    policy_table.rows[0].cells[2].text = "기후여건으로 인한\n불능일 > 8일 적용 경우"
    _shade_header_row(policy_table)

    policy_table.rows[1].cells[0].text = (
        "• 1주 공휴일 수 = 2일 (주7일 - 주5일 = 2일)\n"
        "• 1달 = 4주\n"
        "➡ 1달 공휴일수 = 4주×2일 = 8일"
    )
    policy_table.rows[1].cells[1].text = (
        "• 기후여건 불능일 7일\n"
        "• 법정 근로시간 휴일 8일\n"
        "➡ 비작업일수 8일 적용"
    )
    policy_table.rows[1].cells[2].text = (
        "• 기후여건 불능일 9일\n"
        "• 법정 근로시간 휴일 8일\n"
        "➡ 비작업일수 9일 적용"
    )

    document.add_heading("1.2.3 법정공휴일 수", level=3)
    document.add_paragraph(
        "• 비작업일수 산출을 위한 요소 중 법정공휴일은 공공 기준의 공휴일 데이터를 기준으로 산정하고, "
        "비작업일수에 포함하여 계산한다."
    )

    month_labels = [f"{m}월" for m in range(1, 13)]
    document.add_paragraph("월간 법정공휴일")
    holiday_table = document.add_table(rows=1, cols=14)
    holiday_table.style = "Table Grid"
    header = holiday_table.rows[0].cells
    header[0].text = "년도"
    for idx, label in enumerate(month_labels, start=1):
        header[idx].text = label
    header[13].text = "소계"
    _shade_header_row(holiday_table)

    rows = public_holiday_rows or []
    if rows:
        for entry in rows:
            row = holiday_table.add_row().cells
            row[0].text = str(entry.get("year", "-"))
            monthly = entry.get("monthly", [])
            for idx in range(12):
                value = monthly[idx] if idx < len(monthly) else 0
                row[idx + 1].text = str(value)
            row[13].text = str(entry.get("total", sum(monthly[:12] if monthly else [])))
    else:
        row = holiday_table.add_row().cells
        row[0].text = "데이터 없음"
        for idx in range(1, 14):
            row[idx].text = "-"

    if rows:
        start_year = rows[0]["year"]
        end_year = rows[min(3, len(rows) - 1)]["year"]
        document.add_paragraph(
            f"• 부록 1의 법정 공휴일수 중 당해 공사에 포함되는 {start_year}~{end_year}년 데이터를 적용"
        )
    else:
        document.add_paragraph("• 법정 공휴일 데이터가 없어 표를 산정하지 못함")

    document.add_paragraph(
        "• 당해 공사의 개시일부터 종료일 사이에 포함된 일수를 모두 계산하며, "
        "해당 프로젝트의 조건에 따라 적용범위를 조정한다."
    )
    document.add_paragraph(
        "• 주6일(일요일 휴무) + 필수공휴일(신정, 구정, 근로자의날, 추석, 크리스마스, 선거, 일요일) "
        "기준으로 적용한다."
    )

    document.add_paragraph("※ 공공의 공휴일")
    document.add_paragraph("- 일요일(52일)")
    document.add_paragraph("- 명절(6일): 설 연휴, 추석 연휴 + 대체공휴일 시행")
    document.add_paragraph("- 국경일(4일): 삼일절, 광복절, 개천절, 한글날 + 대체공휴일 시행")
    document.add_paragraph("- 기타(5일): 1월 1일, 5월 5일(대체공휴일 시행), 6월 6일, 부처님 오신 날(음력 4월 8일), 12월 25일")
    document.add_paragraph("- 공직선거법 제34조에 따른 임기만료에 따른 선거의 선거일")
    document.add_paragraph(
        "- 기타 정부에서 수시 지정하는 날 → 대체공휴일 적용. "
        "해당 공휴일이 일요일과 겹치는 경우 대체공휴일, 토요일·일요일과 겹치는 경우 대체공휴일을 적용함"
    )

    document.add_heading("1.3 비작업일수 산정", level=2)
    document.add_heading("1.3.1 개요", level=3)
    document.add_paragraph(
        "• 비작업일수(공사불능일수)의 산정기준은 국토부 고시에 명시된 산출방식에 따라 도출한다."
    )

    p_formula2 = document.add_paragraph()
    p_formula2.add_run(
        "비작업일수 = 기후여건 비작업일수(A) + 해당 월 법정공휴일(B) - 월별 중복일수(C)"
    ).bold = True
    document.add_paragraph("(소수점 첫째자리에서 반올림)")

    document.add_paragraph("• 주 40시간제에 따른 비작업일수 검토")
    document.add_paragraph(
        "- 월간 비작업일수가 주 40시간 근무제에 따른 일수보다 작을 경우에는 "
        "주 40시간 근무제에 따른 비작업일수를 적용한다."
    )
    document.add_paragraph("- 주 40시간제에 따른 비작업일수는 월 8일로 계산한다.")
    document.add_paragraph(
        "- 비작업일수(공사불능일수)는 기후여건과 법정 공휴일이 중복된 경우 1일로 산정한다."
    )

    abc_table = document.add_table(rows=2, cols=3)
    abc_table.style = "Table Grid"
    abc_table.rows[0].cells[0].text = "기후여건으로 인한 비작업일수(A)"
    abc_table.rows[0].cells[1].text = "법정공휴일수(B)"
    abc_table.rows[0].cells[2].text = "중복일수(C)"
    _shade_header_row(abc_table)
    abc_table.rows[1].cells[0].text = (
        "• 기후여건: 기온, 강우, 바람, 적설\n"
        "• 주공정에 영향을 미치는 기상조건 반영\n"
        "• 해당지역 최근 10년간의 기상정보(기상청 기상관측 데이터 적용)\n"
        "• 관계법령과 기준에 공통되는 작업을 제한하는 기상조건 반영"
    )
    abc_table.rows[1].cells[1].text = (
        "• 일요일\n"
        "• 명절: 설연휴, 추석연휴\n"
        "• 국경일: 삼일절, 현충일, 광복절, 개천절, 한글날 등\n"
        "• 기타: 신정, 성탄절, 석가탄신일, 어린이날, 대체휴무 포함"
    )
    abc_table.rows[1].cells[2].text = "• A×B÷달력일수"

    document.add_paragraph("• 아래의 예시와 같이 비작업일수를 산정")
    document.add_paragraph("<예시> 토목공사 비작업일수 산정 (01월 기준)")

    example_table = document.add_table(rows=3, cols=9)
    example_table.style = "Table Grid"
    example_table.rows[0].cells[0].text = "구분"
    example_table.rows[0].cells[1].text = "평균기온"
    example_table.rows[0].cells[2].text = "최고기온"
    example_table.rows[0].cells[3].text = "강우"
    example_table.rows[0].cells[4].text = "미세먼지"
    example_table.rows[0].cells[5].text = "계(A)"
    example_table.rows[0].cells[6].text = "법정공휴일(B)"
    example_table.rows[0].cells[7].text = "중복일수(C)"
    example_table.rows[0].cells[8].text = "비작업일수(A+B-C)"
    _shade_header_row(example_table)

    example_table.rows[1].cells[0].text = "토목공사"
    example_table.rows[1].cells[1].text = "-5℃ 이하"
    example_table.rows[1].cells[2].text = "35℃ 이상"
    example_table.rows[1].cells[3].text = "10mm 이상"
    example_table.rows[1].cells[4].text = "경보 발령시"
    example_table.rows[1].cells[5].text = ""
    example_table.rows[1].cells[6].text = ""
    example_table.rows[1].cells[7].text = ""
    example_table.rows[1].cells[8].text = ""

    example_table.rows[2].cells[0].text = ""
    example_table.rows[2].cells[1].text = "1.6일"
    example_table.rows[2].cells[2].text = "0.0일"
    example_table.rows[2].cells[3].text = "0.3일"
    example_table.rows[2].cells[4].text = "0.1일"
    example_table.rows[2].cells[5].text = "2.0일"
    example_table.rows[2].cells[6].text = "6.7일"
    example_table.rows[2].cells[7].text = "0.4일"
    example_table.rows[2].cells[8].text = "8.3일"

    document.add_paragraph("※ 법정공휴일 = (2022~2024년 1월 공휴일 합산) / 3 = 6.7일")
    document.add_paragraph("※ 중복일수(C) = A × B ÷ 달력일수 = 2.0 × 6.7 ÷ 31 = 0.4일")
    document.add_paragraph("※ 비작업일수 = A + B - C = 2.0 + 6.7 - 0.4 = 8.3일 ≒ 8일(소수점 반올림)")

    document.add_heading("1.3.2 기후 여건으로 인한 비작업일수", level=3)
    document.add_paragraph(
        "• 국토부 고시 부록 3에 제시된 기상 조건별·지역별 비작업일수 조건을 기준으로 검토한다."
    )
    document.add_paragraph(
        "- 해당 공사 지역의 최근 10년 기상관측자료를 반영하며, 당해 공사지역과 가장 인접한 "
        "기상관측자료(예: 인천 기상청 관측자료)를 적용한다."
    )
    document.add_paragraph("- 미세먼지: 미세먼지 경보 발령일 자료를 적용한다.")
    document.add_paragraph(
        "- 관계 법령 및 표준시방서를 검토한 기준으로 적용기준을 검토하고, "
        "주 공정(Critical Path)에 영향을 미치는 기상조건을 반영한다."
    )

    document.add_paragraph("기후 여건으로 인한 비작업일수 적용기준")
    criteria_table = document.add_table(rows=1, cols=4)
    criteria_table.style = "Table Grid"
    criteria_table.rows[0].cells[0].text = "데이터"
    criteria_table.rows[0].cells[1].text = "구분"
    criteria_table.rows[0].cells[2].text = "세부기준"
    criteria_table.rows[0].cells[3].text = "적용여부"
    _shade_header_row(criteria_table)

    rows = climate_criteria_rows or []
    if rows:
        for entry in rows:
            row = criteria_table.add_row().cells
            row[0].text = entry.get("data_source", "경기도 인천시 기상관측 자료")
            row[1].text = entry.get("group", "-")
            row[2].text = entry.get("detail", "-")
            row[3].text = entry.get("apply_text", "-")
    else:
        row = criteria_table.add_row().cells
        row[0].text = "경기도 인천시 기상관측 자료"
        row[1].text = "-"
        row[2].text = "적용기준 데이터 없음"
        row[3].text = "-"

    # Merge repeated labels for visual grouping.
    _merge_same_text_cells(criteria_table, col_idx=0, start_row=1)
    _merge_same_text_cells(criteria_table, col_idx=1, start_row=1)

    document.add_paragraph("")


def build_schedule_report_docx(
    project_name,
    rate_summary,
    ordered_categories,
    grouped_items,
    rate_map,
    region="",
    public_holiday_rows=None,
    climate_criteria_rows=None,
):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for report export. Install 'python-docx' in backend requirements."
        ) from exc

    document = Document()
    _setup_document_defaults(document)

    title = document.add_heading("공사기간 산정 보고서", level=0)
    title.alignment = 1
    _add_duration_analysis_section(
        document,
        public_holiday_rows=public_holiday_rows,
        climate_criteria_rows=climate_criteria_rows,
    )

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.read()
